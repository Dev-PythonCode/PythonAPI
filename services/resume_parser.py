import spacy
import re
import json
from docx import Document
from pathlib import Path


class ResumeParser:
    def __init__(self, spacy_model="en_core_web_sm"):
        self.nlp = spacy.load(spacy_model)
        self.load_skills_data()
        self.setup_matcher()

    def load_skills_data(self):
        skills_path = Path(__file__).parent.parent / "data" / "skills.json"
        with open(skills_path, 'r') as f:
            data = json.load(f)
        self.skills_data = data['skills']
        self.experience_patterns = data['experience_patterns']

    def setup_matcher(self):
        from spacy.matcher import PhraseMatcher
        self.matcher = PhraseMatcher(self.nlp.vocab, attr="LOWER")

        for skill in self.skills_data:
            patterns = [self.nlp.make_doc(skill['name'].lower())]
            for alias in skill.get('aliases', []):
                patterns.append(self.nlp.make_doc(alias.lower()))
            self.matcher.add(skill['name'], patterns)

    def extract_text_from_docx(self, file_path):
        """Extract text from a .docx file"""
        doc = Document(file_path)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text.append(cell.text)
        return '\n'.join(full_text)

    def extract_text_from_pdf(self, file_path):
        """Extract text from a PDF file"""
        from PyPDF2 import PdfReader
        reader = PdfReader(file_path)
        text = ""
        for page in reader.pages:
            text += page.extract_text() + "\n"
        return text

    def extract_skills(self, text):
        """Extract skills from text using SpaCy PhraseMatcher"""
        doc = self.nlp(text.lower())
        matches = self.matcher(doc)

        found_skills = {}
        for match_id, start, end in matches:
            skill_name = self.nlp.vocab.strings[match_id]
            if skill_name not in found_skills:
                found_skills[skill_name] = {
                    'name': skill_name,
                    'count': 0,
                    'contexts': []
                }
            found_skills[skill_name]['count'] += 1

            # Get context around the match
            context_start = max(0, start - 10)
            context_end = min(len(doc), end + 10)
            context = doc[context_start:context_end].text
            found_skills[skill_name]['contexts'].append(context)

        return found_skills

    def extract_experience_years(self, text, skill_name):
        """Try to find years of experience for a specific skill"""
        text_lower = text.lower()
        skill_lower = skill_name.lower()

        # Find skill mentions and look for nearby year patterns
        patterns = [
            rf'{skill_lower}\s*[-â€“:]\s*([0-9]+(?:\.[0-9]+)?)\s*(?:years?|yrs?)',
            rf'([0-9]+(?:\.[0-9]+)?)\s*(?:years?|yrs?)\s*(?:of\s*)?{skill_lower}',
            rf'{skill_lower}.?([0-9]+(?:\.[0-9]+)?)\s(?:years?|yrs?)',
        ]

        for pattern in patterns:
            match = re.search(pattern, text_lower)
            if match:
                return float(match.group(1))

        return None

    def estimate_proficiency(self, years, count):
        """Estimate proficiency level based on years and mention count"""
        if years:
            if years >= 5:
                return "Expert"
            elif years >= 3:
                return "Advanced"
            elif years >= 1:
                return "Intermediate"
            else:
                return "Beginner"
        else:
            # Estimate based on mention count
            if count >= 5:
                return "Advanced"
            elif count >= 3:
                return "Intermediate"
            else:
                return "Beginner"

    def calculate_confidence(self, years, count):
        """Calculate confidence score for extracted skill"""
        base_confidence = 0.5

        if years:
            base_confidence += 0.3

        if count >= 3:
            base_confidence += 0.2
        elif count >= 2:
            base_confidence += 0.1

        return min(base_confidence, 1.0)

    def parse_resume(self, file_path):
        """Main method to parse a resume and extract skills"""
        # Determine file type and extract text
        file_path = Path(file_path)
        if file_path.suffix.lower() == '.docx':
            text = self.extract_text_from_docx(file_path)
        elif file_path.suffix.lower() == '.pdf':
            text = self.extract_text_from_pdf(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_path.suffix}")

        # Extract skills
        found_skills = self.extract_skills(text)

        # Build results with experience and proficiency
        results = []
        for skill_name, skill_data in found_skills.items():
            years = self.extract_experience_years(text, skill_name)
            proficiency = self.estimate_proficiency(years, skill_data['count'])
            confidence = self.calculate_confidence(years, skill_data['count'])

            results.append({
                'skill_name': skill_name,
                'years_of_experience': years,
                'proficiency_level': proficiency,
                'confidence_score': confidence,
                'mention_count': skill_data['count']
            })

        # Sort by confidence
        results.sort(key=lambda x: x['confidence_score'], reverse=True)

        return {
            'extracted_skills': results,
            'total_skills_found': len(results),
            'raw_text_length': len(text)
        }

    def load_skills_data(self):
        skills_path = Path(__file__).parent.parent / "data" / "skills.json"
        with open(skills_path, 'r') as f:
            data = json.load(f)
        self.skills_data = data['skills']
        self.locations = data.get('locations', [])
        self.experience_patterns = data.get('experience_patterns', [])
